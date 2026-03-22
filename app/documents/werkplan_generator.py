"""Werkplan generator — Word (.docx) output in 3D-Drilling standaardformaat.

Genereert een compleet werkplan voor een gestuurde boring op basis van
Order- en Boring-data uit de database. Het formaat volgt het standaard
werkplan template van GestuurdeBoring Tekening.nl.
"""
from __future__ import annotations

import io
import math
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from app.order.models import Boring, Order, WerkplanAfbeelding
from app.order.klantcodes import get_klant_naam


# ── Helpers ────────────────────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex: str):
    """Achtergrondkleur voor een tabelcel instellen."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def _add_styled_paragraph(doc, text: str, style: str = "Normal",
                           bold: bool = False, size: int = 11,
                           alignment=None, space_after: int | None = None,
                           color: RGBColor | None = None):
    """Voeg een opgemaakte paragraaf toe."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def _add_heading(doc, text: str, level: int = 1):
    """Voeg een heading toe in huisstijl (donkerblauw)."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # Donkerblauw
        run.font.name = "Calibri"
    return h


def _format_datum(d: date | None = None) -> str:
    """Datum formatteren als '17 februari 2026'."""
    if d is None:
        d = date.today()
    maanden = [
        "", "januari", "februari", "maart", "april", "mei", "juni",
        "juli", "augustus", "september", "oktober", "november", "december",
    ]
    return f"{d.day} {maanden[d.month]} {d.year}"


def _aantal_buizen_tekst(boring: Boring) -> str:
    """Genereer buisspecificatie tekst, bv '1x buis Ø110 SDR11 (wanddikte 10mm)'."""
    wanddikte = boring.dn_effectief
    return (
        f"1x buis Ø{boring.De_mm:.0f} SDR{boring.SDR} "
        f"(wanddikte {wanddikte:.0f}mm)"
    )


def _add_image(doc, bestandspad: str, bijschrift: str = "",
                width_cm: float = 12.0):
    """Voeg een afbeelding toe aan het document met optioneel bijschrift."""
    path = Path(bestandspad)
    if not path.exists():
        _add_styled_paragraph(
            doc, f"[Afbeelding niet gevonden: {path.name}]",
            color=RGBColor(0xCC, 0x00, 0x00), space_after=6,
        )
        return
    try:
        doc.add_picture(str(path), width=Cm(width_cm))
        if bijschrift:
            p = doc.add_paragraph()
            run = p.add_run(bijschrift)
            run.italic = True
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(6)
    except Exception:
        _add_styled_paragraph(
            doc, f"[Kon afbeelding niet laden: {path.name}]",
            color=RGBColor(0xCC, 0x00, 0x00), space_after=6,
        )


def _generate_werkplan_kaart(boring) -> str | None:
    """Genereer een OSM kaart als tijdelijk JPG bestand voor het werkplan.

    Returns: pad naar tijdelijk bestand, of None.
    """
    trace_punten = [p for p in boring.trace_punten if getattr(p, 'variant', 0) == 0]
    if len(trace_punten) < 2:
        return None
    try:
        import tempfile
        import math as _m
        from app.geo.coords import rd_to_wgs84
        from app.documents.pdf_generator import _fetch_map_image_b64

        xs = [p.RD_x for p in trace_punten]
        ys = [p.RD_y for p in trace_punten]
        cx = (min(xs) + max(xs)) / 2
        cy = (min(ys) + max(ys)) / 2
        lat_c, lon_c = rd_to_wgs84(cx, cy)

        trace_span = max(max(xs) - min(xs), max(ys) - min(ys))
        zm = 19 if trace_span < 500 else 18

        b64 = _fetch_map_image_b64(lat_c, lon_c, zoom=zm, tiles_x=7, tiles_y=4)
        if not b64:
            return None

        import base64, io
        from PIL import Image, ImageDraw, ImageFont

        img_bytes = base64.b64decode(b64.split(",", 1)[1])
        img = Image.open(io.BytesIO(img_bytes))
        draw = ImageDraw.Draw(img)
        img_w, img_h = img.size
        mpp = 40075016.686 * _m.cos(_m.radians(lat_c)) / (256.0 * (2 ** zm))

        def rd_to_px(rx, ry):
            return int(img_w/2 + (rx - cx)/mpp), int(img_h/2 - (ry - cy)/mpp)

        # Tracélijn
        tpx = [rd_to_px(p.RD_x, p.RD_y) for p in trace_punten]
        if len(tpx) >= 2:
            draw.line(tpx, fill=(204, 0, 0), width=5)
        try:
            font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 16)
        except Exception:
            font = ImageFont.load_default()
        for i, p in enumerate(trace_punten):
            px, py = tpx[i]
            draw.ellipse([px-5, py-5, px+5, py+5], fill=(204, 0, 0), outline=(255, 255, 255))
            if p.label:
                draw.rectangle([px+8, py-18, px+8+len(p.label)*10, py], fill=(255, 255, 255))
                draw.text((px+8, py-18), p.label, fill=(204, 0, 0), font=font)

        # Crop rond tracé
        margin_px = 120
        trace_xs = [t[0] for t in tpx]
        trace_ys = [t[1] for t in tpx]
        cl = max(0, min(trace_xs) - margin_px)
        ct = max(0, min(trace_ys) - margin_px)
        cr = min(img_w, max(trace_xs) + margin_px)
        cb = min(img_h, max(trace_ys) + margin_px)
        if (cb - ct) < (cr - cl) * 0.4:
            extra = int(((cr - cl) * 0.4 - (cb - ct)) / 2)
            ct = max(0, ct - extra)
            cb = min(img_h, cb + extra)
        img = img.crop((cl, ct, cr, cb))

        f = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
        img.save(f, format="JPEG", quality=88)
        f.flush()
        return f.name
    except Exception:
        return None


def _get_afbeeldingen(boring: Boring, categorie: str) -> list:
    """Haal afbeeldingen op voor een boring per categorie."""
    return [
        a for a in (boring.werkplan_afbeeldingen or [])
        if a.categorie == categorie
    ]


def _medium_tekst(boring: Boring) -> str:
    """Medium omschrijving voor buisspecificatie."""
    medium_map = {
        "Drukloos": "Geen (mantelbuis voor Laagspanningskabels)",
        "Druk": "Drukleiding",
        "Gas": "Gasleiding",
        "Middenspanning": "Geen (mantelbuis voor middenspanning)",
        "Water": "Waterleiding",
        "Telecom": "Geen (mantelbuis voor telecom)",
    }
    return medium_map.get(boring.medium, boring.medium or "Geen")


def _ckb_categorie(boring: Boring) -> tuple[str, str]:
    """Bepaal CKB-categorie op basis van berekende intrekkracht.

    ST-A (<9T) · ST-B (10–39T) · ST-C (40–149T) · ST-D (>150T)
    """
    if boring.berekening and boring.berekening.Ttot_N:
        ton = abs(boring.berekening.Ttot_N) / 9806.65  # N naar metrische ton
        if ton < 9:
            return "ST-A", "Kleine gestuurde boringen (maximaal 9 ton)"
        elif ton <= 39:
            return "ST-B", "Grote gestuurde boringen (van 9 tot 39 ton)"
        elif ton <= 149:
            return "ST-C", "Zeer grote gestuurde boringen (van 40 tot 149 ton)"
        else:
            return "ST-D", "Extra grote gestuurde boringen (meer dan 150 ton)"
    return "ST-A", "Kleine gestuurde boringen (maximaal 9 ton)"


def _generate_inleiding_template(doc, klant_naam: str, hoofdaannemer: str):
    """Fallback template inleiding als er geen AI of eigen tekst is."""
    klant_tekst = klant_naam if klant_naam else "[Uitvoerend bedrijf]"
    _add_styled_paragraph(
        doc,
        f"Om aan te tonen dat de boring veilig en volgens voorschriften uitgevoerd kan worden, "
        f"is dit werkplan opgesteld.",
        space_after=6,
    )
    _add_styled_paragraph(
        doc,
        f"{klant_tekst} voert de gestuurde boring uit"
        + (f" in opdracht van {hoofdaannemer}." if hoofdaannemer else "."),
        space_after=6,
    )
    _add_styled_paragraph(
        doc,
        f"In opdracht van {klant_tekst} heeft GestuurdeBoringTekening.nl "
        f"dit werkplan opgesteld.",
        space_after=6,
    )


def _generate_kwel_template(doc, kwel_gebied: bool):
    """Fallback kwel-template tekst."""
    if kwel_gebied:
        _add_styled_paragraph(
            doc,
            "[Kwelmaatregelen beschrijven - het betreft hier een kwetsbaar kwelgebied]",
            space_after=12,
            color=RGBColor(0xCC, 0x00, 0x00),
        )
    else:
        _add_styled_paragraph(
            doc,
            "Er is hier geen sprake van een kwetsbaar kwelgebied.",
            space_after=6,
        )
        _add_styled_paragraph(
            doc,
            "Er wordt met deze boring geen waterkering gekruist en geen boezemwater onderkruist.",
            space_after=6,
        )
        _add_styled_paragraph(
            doc,
            "Conclusie: Er zijn geen aanvullende maatregelen nodig om problemen met kwel voor te zijn.",
            space_after=12,
        )


# ── Hoofdgenerator ─────────────────────────────────────────────────────────

def generate_werkplan(order: Order, boring: Boring,
                      auteur: str = "Martien Luijben",
                      hoofdaannemer: str = "",
                      opdrachtgever_naam: str = "",
                      inleiding_tekst: str = "",
                      kwel_gebied: bool = False,
                      revisie: int = 0,
                      revisie_omschrijving: str = "Vergunningsaanvraag",
                      gebruik_ai: bool = False,
                      db=None,
                      ) -> bytes:
    """Genereer een werkplan als Word-document (.docx bytes).

    Args:
        order: Order object met alle relaties geladen
        boring: Boring object met trace_punten, doorsneden, berekening
        auteur: Naam van de auteur (default: Martien Luijben)
        hoofdaannemer: Naam hoofdaannemer
        opdrachtgever_naam: Naam opdrachtgever (logo-tekst)
        inleiding_tekst: Vrije tekst voor de inleiding
        kwel_gebied: Of het project in een kwelgebied ligt
        revisie: Revisienummer
        revisie_omschrijving: Omschrijving bij revisie

    Returns:
        bytes: Het werkplan als .docx bestand
    """
    doc = Document()

    # ── Document-instellingen ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Klant-info afleiden
    klant_naam = get_klant_naam(order.klantcode) if order.klantcode else ""
    opdracht_naam = opdrachtgever_naam or order.opdrachtgever or ""
    locatie = order.locatie or ""
    vandaag = _format_datum()

    # ── VOORPAGINA ──────────────────────────────────────────────────────

    # Lege regels voor spacing bovenaan
    for _ in range(2):
        doc.add_paragraph()

    # Titel
    _add_styled_paragraph(
        doc, "Werkplan Gestuurde Boring",
        bold=True, size=26,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6,
    )

    # Locatie als subtitel
    if locatie:
        _add_styled_paragraph(
            doc, locatie,
            bold=True, size=22,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12,
        )

    # Projectomschrijving
    if boring.naam:
        _add_styled_paragraph(
            doc, boring.naam,
            bold=True, size=14,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24,
        )

    # Opdrachtgever
    _add_styled_paragraph(doc, "Opdrachtgever:", size=11, space_after=2)
    _add_styled_paragraph(
        doc, opdracht_naam or "[Opdrachtgever]",
        bold=True, size=14, space_after=12,
    )

    # Hoofdaannemer
    if hoofdaannemer:
        _add_styled_paragraph(doc, "Hoofdaannemer:", size=11, space_after=2)
        _add_styled_paragraph(
            doc, hoofdaannemer,
            bold=True, size=14, space_after=12,
        )

    # Spacing
    for _ in range(3):
        doc.add_paragraph()

    # Opdrachtnummer + Auteur
    p = doc.add_paragraph()
    run = p.add_run(f"Opdrachtnummer:\t{order.ordernummer}")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    p = doc.add_paragraph()
    run = p.add_run(f"Auteur:\t\t{auteur}")
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    doc.add_paragraph()

    # GestuurdeBoring Tekening.nl tekst
    _add_styled_paragraph(
        doc, "GestuurdeBoringTekening.nl",
        bold=True, size=16,
        color=RGBColor(0x2E, 0x74, 0xB5),
        space_after=12,
    )

    # Revisietabel
    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header rij
    headers = ["Revisie", "Datum", "Omschrijving", "Voorbereid", "Goedgekeurd"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Calibri"

    # Data rij
    akkoord = order.akkoord_contact or ""
    data = [str(revisie), vandaag, revisie_omschrijving, auteur.split()[0] + "." + auteur.split()[-1][0] if " " in auteur else auteur, akkoord]
    for i, val in enumerate(data):
        cell = table.rows[1].cells[i]
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = "Calibri"

    # Pagina-einde na voorpagina
    doc.add_page_break()

    # ── INHOUDSOPGAVE (placeholder) ─────────────────────────────────────

    _add_heading(doc, "Inhoud", level=1)
    inhoud_items = [
        ("1.", "Inleiding", 2),
        ("2.", "Voorbereiding", 2),
        ("2.1.", "De locatie", 3),
        ("2.2.", "Historie", 3),
        ("2.3.", "Infrastructuur", 3),
        ("3.", "Ontwerp van de gestuurde boring en het gegraven tracé", 2),
        ("3.1.", "Het boorprofiel en gegraven tracé", 3),
        ("3.2.", "Metingen en registraties", 3),
        ("3.2.1", "Het boorkop volgsysteem", 4),
        ("3.2.2", "Voorbereiding Bentoniet", 4),
        ("3.2.3", "Boorspoelparameters", 4),
        ("3.2.4", "Rijksdriehoekcoördinaten", 4),
        ("4.", "Sterkte- Intrek- en Boorspoeldrukberekeningen", 2),
        ("4.1.", "Sterkteberekeningen", 3),
        ("4.2.", "Berekening benodigde intrekkracht", 3),
        ("4.3.", "Boorspoeldrukberekeningen", 3),
        ("5.", "Kwel", 2),
        ("5.1.", "Oorzaak", 3),
        ("5.2.", "Risico's en maatregelen", 3),
        ("5.3.", "Maatregelen op dit project", 3),
        ("6.", "Uitvoering", 2),
        ("6.1", "In te zetten boormachine", 3),
        ("6.1.1", "Details boormachine", 4),
        ("6.2", "Personeel", 3),
        ("6.3", "Tijdsplanning", 3),
        ("6.4", "Afwijkingen", 3),
        ("6.5", "Inspectie na intrekken", 3),
    ]

    for nr, titel, indent_level in inhoud_items:
        p = doc.add_paragraph()
        indent_cm = (indent_level - 2) * 0.75
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{nr}\t{titel}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        if indent_level == 2:
            run.bold = True

    # Bijlagen
    bijlagen = [
        "Bijlage A: Tekening van de gestuurde boring",
        "Bijlage B: Berekeningen",
        "Bijlage C: Geotechnische bodemgegevens",
        "Bijlage D: Procesbeschrijving techniek van gestuurd boren",
        "Bijlage E: De boormachine",
        "Bijlage F: Het boorkop-volgsysteem",
        "Bijlage G: Bentoniet",
    ]
    doc.add_paragraph()
    for bijlage in bijlagen:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(bijlage)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    doc.add_page_break()

    # ── HOOFDSTUK 1: INLEIDING ──────────────────────────────────────────

    _add_heading(doc, "1.  Inleiding", level=1)

    if inleiding_tekst:
        # Gebruiker heeft eigen tekst opgegeven
        for alinea in inleiding_tekst.split("\n\n"):
            if alinea.strip():
                _add_styled_paragraph(doc, alinea.strip(), space_after=6)
    elif gebruik_ai:
        # AI-gegenereerde inleiding
        try:
            from app.ai_assist.werkplan_teksten import genereer_inleiding
            ai_tekst = genereer_inleiding(order, boring, hoofdaannemer)
            if ai_tekst:
                for alinea in ai_tekst.split("\n\n"):
                    if alinea.strip():
                        _add_styled_paragraph(doc, alinea.strip(), space_after=6)
            else:
                raise ValueError("Lege AI response")
        except Exception:
            # Fallback naar template
            _generate_inleiding_template(doc, klant_naam, hoofdaannemer)
    else:
        _generate_inleiding_template(doc, klant_naam, hoofdaannemer)

    doc.add_page_break()

    # ── HOOFDSTUK 2: VOORBEREIDING ──────────────────────────────────────

    _add_heading(doc, "2.  Voorbereiding", level=1)
    _add_styled_paragraph(
        doc,
        "Alvorens te starten met het ontwerpen van de gestuurde boring, "
        "worden diverse omgevingsfactoren in kaart gebracht:",
        space_after=12,
    )

    # 2.1 De locatie
    _add_heading(doc, "2.1. De locatie", level=2)

    # AI locatiebeschrijving of placeholder
    if gebruik_ai:
        try:
            from app.ai_assist.werkplan_teksten import genereer_locatie_beschrijving
            loc_tekst = genereer_locatie_beschrijving(order, boring)
            if loc_tekst:
                _add_styled_paragraph(doc, loc_tekst, space_after=6)
            else:
                _add_styled_paragraph(doc, f"[Beschrijving van de locatie: {locatie}]", space_after=6)
        except Exception:
            _add_styled_paragraph(doc, f"[Beschrijving van de locatie: {locatie}]", space_after=6)
    else:
        _add_styled_paragraph(doc, f"[Beschrijving van de locatie: {locatie}]", space_after=6)

    # Luchtfoto: eerst uit DB, dan auto-genereren uit tracé
    luchtfotos = _get_afbeeldingen(boring, "luchtfoto")
    if luchtfotos:
        for foto in luchtfotos:
            _add_image(doc, foto.bestandspad, foto.bijschrift or "")
    else:
        kaart_path = _generate_werkplan_kaart(boring)
        if kaart_path:
            _add_image(doc, kaart_path, "Situatiekaart tracé (automatisch gegenereerd)", width_cm=14.0)
        else:
            _add_styled_paragraph(
                doc,
                "[Voeg hier een luchtfoto of Google Maps screenshot in]",
                space_after=12,
                color=RGBColor(0xCC, 0x00, 0x00),
            )

    # 2.2 Historie
    _add_heading(doc, "2.2.  Historie", level=2)
    _add_styled_paragraph(
        doc,
        "Via topotijdreis.nl is nagegaan of op de locatie wegen of bouwwerken gelegen hebben, "
        "om in een vroeg stadium te bepalen of er kans is oude fundaties tegen te komen, "
        "die een aanpassing van het beoogde tracé tot gevolg kunnen hebben.",
        space_after=6,
    )

    # Topotijdreis screenshots uit DB
    topo_fotos = sorted(
        [a for a in (boring.werkplan_afbeeldingen or [])
         if a.categorie.startswith("topotijdreis")],
        key=lambda a: a.volgorde,
    )
    if topo_fotos:
        for foto in topo_fotos:
            _add_image(doc, foto.bestandspad, foto.bijschrift or "", width_cm=5.0)
        doc.add_paragraph()
    else:
        _add_styled_paragraph(
            doc,
            "[Voeg hier 3 screenshots van topotijdreis.nl in]",
            space_after=6,
            color=RGBColor(0xCC, 0x00, 0x00),
        )

    _add_styled_paragraph(
        doc,
        "Conclusie: Vanuit de historie is geen aanleiding gevonden om problemen met funderingsresten "
        "te verwachten.",
        space_after=12,
    )

    # 2.3 Infrastructuur
    _add_heading(doc, "2.3.  Infrastructuur", level=2)
    _add_styled_paragraph(
        doc,
        "Aanwezige Kabels in Leidingen in de ondergrond: Middels KLIC oriëntatiemelding via kadaster.nl "
        "zijn de bestaande Kabels en Leidingen in kaart gebracht en is het tracé van de boring zodanig "
        "ontworpen dat de kans op schade door graafwerkzaamheden, het boren of het ruimen tot een "
        "minimum beperkt worden. Zie voor het ontwerp van de gestuurde boring bijlage A.",
        space_after=6,
    )

    _add_styled_paragraph(
        doc,
        "Aanlegvoorschriften: De diverse grondeigenaren, Kabels en Leidingbeheerders, Waterschappen, "
        "Provincies en Rijkswaterstaat hebben voorschriften opgesteld waaraan moet worden voldaan bij "
        "werkzaamheden in de buurt van hun assets. Deze voorschriften zijn in kaart gebracht en in het "
        "ontwerp van het boortracé, de tekening en/of de berekening verwerkt.",
        space_after=6,
    )

    _add_styled_paragraph(doc, "KABELS & LEIDINGEN:", bold=True, space_after=4)

    # ── KLIC data uit platform ──
    klic_data_gevuld = False
    if db is not None:
        from app.order.models import KLICLeiding, KLICUpload
        laatste_upload = (
            db.query(KLICUpload)
            .filter_by(order_id=order.id, verwerkt=True)
            .order_by(KLICUpload.upload_datum.desc())
            .first()
        )
        if laatste_upload:
            leidingen = db.query(KLICLeiding).filter_by(klic_upload_id=laatste_upload.id).all()
            if leidingen:
                klic_data_gevuld = True
                # Samenvatting per beheerder
                _add_styled_paragraph(
                    doc,
                    f"Op basis van KLIC melding {laatste_upload.meldingnummer or ''} "
                    f"zijn {len(leidingen)} kabels en leidingen van "
                    f"{laatste_upload.aantal_beheerders or '?'} beheerders in kaart gebracht:",
                    space_after=4,
                )
                # Tabel: beheerder | type | aantal
                from collections import defaultdict
                agg = defaultdict(lambda: {"aantal": 0, "types": set()})
                for l in leidingen:
                    key = l.beheerder or "Onbekend"
                    agg[key]["aantal"] += 1
                    if l.leidingtype:
                        agg[key]["types"].add(l.leidingtype)
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                hdr = table.rows[0].cells
                hdr[0].text = "Beheerder"
                hdr[1].text = "Leidingtype"
                hdr[2].text = "Aantal"
                for cell in hdr:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(9)
                for beh in sorted(agg.keys()):
                    row = table.add_row().cells
                    row[0].text = beh
                    row[1].text = ", ".join(sorted(agg[beh]["types"]))[:60]
                    row[2].text = str(agg[beh]["aantal"])
                    for cell in row:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9)
                doc.add_paragraph()

                # Sleufloze leidingen waarschuwing
                sleufloze = [l for l in leidingen if l.sleufloze_techniek]
                if sleufloze:
                    _add_styled_paragraph(
                        doc,
                        f"Let op: {len(sleufloze)} leiding(en) zijn gedetecteerd als sleufloze techniek "
                        f"(HDD/boogzinker). Deze leidingen liggen mogelijk dieper dan verwacht.",
                        bold=True, space_after=4,
                        color=RGBColor(0xCC, 0x66, 0x00),
                    )

                # Diepte waarschuwing
                met_diepte = sum(1 for l in leidingen if l.diepte_m is not None)
                if met_diepte == 0:
                    _add_styled_paragraph(
                        doc,
                        "Opmerking: Geen van de leidingen in de KLIC heeft betrouwbare dieptegegevens. "
                        "Dieptes dienen ter plaatse te worden gecontroleerd conform CROW 96b.",
                        space_after=6,
                        color=RGBColor(0xCC, 0x00, 0x00),
                    )

    # KLIC screenshot/kaart
    klic_fotos = _get_afbeeldingen(boring, "klic")
    if klic_fotos:
        for foto in klic_fotos:
            _add_image(doc, foto.bestandspad, foto.bijschrift or "Screenshot van de KLIC")
    else:
        # Auto-genereer kaart als situatiekaart beschikbaar
        kaart_path = _generate_werkplan_kaart(boring)
        if kaart_path:
            _add_image(doc, kaart_path, "Situatie K&L (automatisch gegenereerd)", width_cm=14.0)
        elif not klic_data_gevuld:
            _add_styled_paragraph(
                doc,
                "[Voeg hier een screenshot van de KLIC viewer in + analyse]",
                space_after=6,
                color=RGBColor(0xCC, 0x00, 0x00),
        )

    # EV-partijen uit platform
    if order.ev_partijen:
        _add_styled_paragraph(doc, "EIS VOORZORGSMAATREGEL (EV):", bold=True, space_after=4)
        _add_styled_paragraph(
            doc,
            "De volgende netbeheerders hebben een Eis Voorzorgsmaatregel (EV) op hun leidingen. "
            "Vóór aanvang van werkzaamheden dient contact te worden opgenomen:",
            space_after=4,
        )
        for ev in order.ev_partijen:
            _add_styled_paragraph(doc, f"  • {ev.naam}", space_after=2)
        doc.add_paragraph()

    _add_styled_paragraph(
        doc,
        "De gebruikelijke voorzichtigheid en werkwijze volgens CROW500 zijn hier voldoende "
        "om het risico op schade uit te sluiten.",
        space_after=12,
    )

    # Waterschap info
    if order.waterkering_url:
        _add_styled_paragraph(doc, "WATERSCHAP:", bold=True, space_after=4)
        _add_styled_paragraph(
            doc,
            "Het tracé ligt binnen het beheergebied van het waterschap. De waterkering- en "
            "watergang-voorschriften zijn gecontroleerd en in het ontwerp verwerkt.",
            space_after=6,
        )

    # RWS paragraaf als vergunning = R
    if order.vergunning == "R":
        _add_styled_paragraph(doc, "RIJKSWATERSTAAT", bold=True, space_after=4)
        _add_styled_paragraph(
            doc,
            "Graafwerkzaamheden en boringen binnen de beheerzone van Rijkswaterstaat dienen vergund te "
            "worden. Om aan te tonen dat de gestuurde boring veilig uitgevoerd kan worden, vraagt "
            "Rijkswaterstaat bij de vergunningsaanvraag om een boorplan inclusief Sterkte- Intrek- en "
            "Boorspoeldrukberekeningen, dat voldoet aan de Richtlijn Boortechnieken (laatste versie: "
            "Juni 2019-v1.0).",
            space_after=6,
        )
        # RWS screenshot uit DB
        rws_fotos = _get_afbeeldingen(boring, "rws")
        if rws_fotos:
            for foto in rws_fotos:
                _add_image(doc, foto.bestandspad, foto.bijschrift or "Screenshot van de beheergrenzen op rijkswaterstaat.nl")
        else:
            _add_styled_paragraph(
                doc,
                "[Voeg hier een screenshot van de beheergrenzen op rijkswaterstaat.nl in]",
                space_after=6,
                color=RGBColor(0xCC, 0x00, 0x00),
            )
        _add_styled_paragraph(
            doc,
            "De richtlijn schrijft voor dat in- en uittredepunten zodanig dienen worden gekozen "
            "dat de stabiliteit van de aardebaan gewaarborgd blijft.",
            space_after=6,
        )
        _add_styled_paragraph(
            doc,
            "Paragraaf 2.3.4 Kruising met een rijkswaterstaatswerk. Deze paragraaf stelt dat een "
            "rijkswaterstaatswerk in principe loodrecht op de lengte richting gekruist dient te worden. "
            "Hieraan voldoet het ontwerp.",
            space_after=6,
        )
        _add_styled_paragraph(
            doc,
            "Paragraaf 2.3.8 Zettingen en monitoring. Deze paragraaf schrijft voor dat indien zettingen "
            "> 10 mm verwacht kunnen worden, er een zettingsberekening dient te worden uitgevoerd.",
            space_after=6,
        )
        _add_styled_paragraph(
            doc,
            "[Conclusie over zettingen invullen]",
            space_after=6,
            color=RGBColor(0xCC, 0x00, 0x00),
        )
        _add_styled_paragraph(
            doc,
            "Voor de overige paragraven wordt verwezen naar de Sigma Sterkte-, Intrek- en "
            "Boorspoeldrukberekeningen om aan te tonen dat hieraan wordt voldaan.",
            space_after=12,
        )

    doc.add_page_break()

    # ── HOOFDSTUK 3: ONTWERP ────────────────────────────────────────────

    _add_heading(doc, "3.  Ontwerp van de gestuurde boring en het gegraven tracé", level=1)
    _add_styled_paragraph(
        doc,
        "De grove locatie voor de gestuurde boring is vooraf bepaald door "
        f"{opdracht_naam or '[opdrachtgever]'}; de exacte positie van in- en uittrede zijn op basis "
        "van de aanwezige K&L en schouw middels recente beelden van Google Maps vastgesteld.",
        space_after=12,
    )

    # Buisspecificaties
    _add_styled_paragraph(doc, "Specificaties van de aan te leggen buis:", space_after=4)

    spec_table = doc.add_table(rows=5, cols=2)
    spec_data = [
        ("Afmeting:", _aantal_buizen_tekst(boring)),
        ("Materiaal:", boring.materiaal or "PE100"),
        ("Medium:", _medium_tekst(boring)),
        ("Ontwerpdruk:", "ambient"),
        ("Diameter boorgat:", f"Ø{boring.Dg_mm:.0f} mm" if boring.Dg_mm else "[in te vullen]"),
    ]
    for i, (label, val) in enumerate(spec_data):
        spec_table.rows[i].cells[0].text = label
        spec_table.rows[i].cells[1].text = val
        for cell in spec_table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = "Calibri"

    doc.add_paragraph()

    # 3.1 Boorprofiel en gegraven tracé
    _add_heading(doc, "3.1. Het boorprofiel en gegraven tracé", level=2)
    _add_styled_paragraph(
        doc,
        "In het ontwerp van de boring zijn de boorlijn, het maaiveld volgens AHN4 en de kabels en "
        "leidingen volgens de KLIC melding weergegeven in bovenaanzicht en in lengteprofiel.",
        space_after=6,
    )
    _add_styled_paragraph(
        doc,
        "In het ontwerp van het tracé staan zowel de gestuurde boring alsook het gegraven tracé "
        "weergegeven.",
        space_after=6,
    )
    _add_styled_paragraph(doc, "Zie bijlage A voor de tekeningen", space_after=12)

    # 3.2 Metingen en registraties
    _add_heading(doc, "3.2. Metingen en registraties", level=2)
    _add_styled_paragraph(
        doc,
        "Om de boring volgens ontwerp uit te voeren, worden tijdens uitvoering diverse metingen "
        "gemonitord en geregistreerd.",
        space_after=12,
    )

    # 3.2.1 Boorkop volgsysteem
    _add_heading(doc, "3.2.1   Het boorkop volgsysteem", level=3)
    _add_styled_paragraph(
        doc,
        "De diepte, rotatie, hellingshoek en richting van de boorkop worden tijdens het uitvoering "
        "van de pilotboring continue gemonitord met een meetsysteem en teruggekoppeld naar de "
        "boormeester. Deze informatie gebruikt de boormeester om bij te sturen indien nodig en "
        "wordt genoteerd en/of opgeslagen voor latere verwerking in de As-Built tekening van de boring.",
        space_after=6,
    )
    _add_styled_paragraph(
        doc,
        "Tevens worden de temperatuur van de boorkop en de batterijstatus gemonitord.",
        space_after=6,
    )
    _add_styled_paragraph(
        doc,
        "Gegevens van het boorkop volgsysteem zijn terug te vinden in bijlage F.",
        space_after=12,
    )

    # 3.2.2 Voorbereiding Bentoniet
    _add_heading(doc, "3.2.2   Voorbereiding Bentoniet", level=3)
    _add_styled_paragraph(
        doc,
        "Voor het aanmaken van bentoniet wordt gebruik gemaakt van oppervlaktewater. Dit wordt "
        "gecontroleerd op onderstaande parameters. Indien nodig wordt met toevoegmiddelen (polymeren) "
        "de parameter aangepast, of water van elders betrokken met een tankwagen:",
        space_after=6,
    )
    _add_styled_paragraph(doc, "Zuurtegraad, ofwel PH waarde tussen 7 en 9,5", space_after=2)
    _add_styled_paragraph(doc, "Hardheid, ofwel CA < 100 ppm", space_after=2)
    _add_styled_paragraph(doc, "Geleidbaarheid, ofwel Cl < 1000 μS", space_after=6)
    _add_styled_paragraph(
        doc,
        "De datasheet van de te gebruiken bentoniet zijn terug te vinden in bijlage G.",
        space_after=12,
    )

    # 3.2.3 Boorspoelparameters
    _add_heading(doc, "3.2.3   Boorspoelparameters", level=3)
    _add_styled_paragraph(
        doc,
        "Metingen van de boorspoeldruk en boorspoeldebiet worden ingesteld en afgelezen op de "
        "boormachine. Deze worden handmatig geregistreerd tijdens de uitvoering.",
        space_after=12,
    )

    # 3.2.4 RD-coördinaten
    _add_heading(doc, "3.2.4   Rijksdriehoekcoördinaten", level=3)
    _add_styled_paragraph(
        doc,
        "Na uitvoering van de boring worden in- en uittredepunten en eventueel tussenliggende punten "
        "– welke tijdens de pilotfase met spuitbus op het maaiveld gemarkeerd zijn – met GPS meetstok "
        "ingemeten voor verwerking in As-Built tekening. Indien met een gyroscoop gemeten wordt, zal "
        "via dit systeem ook de afgelegde boorgang geregistreerd worden en verwerkt worden in "
        "As-Built tekening.",
        space_after=6,
    )

    # RD-coördinaten tabel als er tracepunten zijn
    if boring.trace_punten:
        rd_table = doc.add_table(rows=len(boring.trace_punten) + 1, cols=3)
        rd_table.style = "Table Grid"
        rd_headers = ["Punt", "RD X", "RD Y"]
        for i, h in enumerate(rd_headers):
            cell = rd_table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"

        for j, punt in enumerate(boring.trace_punten):
            row = rd_table.rows[j + 1]
            row.cells[0].text = punt.label or punt.type
            row.cells[1].text = f"{punt.RD_x:.2f}"
            row.cells[2].text = f"{punt.RD_y:.2f}"
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.name = "Calibri"
        doc.add_paragraph()

    doc.add_page_break()

    # ── HOOFDSTUK 4: BEREKENINGEN ───────────────────────────────────────

    _add_heading(doc, "4.  Sterkte- Intrek- en Boorspoeldrukberekeningen", level=1)
    _add_styled_paragraph(
        doc,
        "De berekeningen aan het ontwerp van deze gestuurde boring zijn conform NEN 3650 en NEN 3651 "
        "en zijn gemaakt met behulp van het pakket Sigma 2024 1.5 van Adviesbureau Schrijvers. De "
        "volledige berekeningen zijn te vinden in bijlage B.",
        space_after=6,
    )
    _add_styled_paragraph(
        doc,
        "Voor de geotechnische gegevens van de ondergrond is uitgegaan van twee sonderingen aan beide "
        "zijden van de gestuurde boring welke terug te vinden zijn in bijlage C.",
        space_after=6,
    )

    # Doorsneden uit platform
    if boring.doorsneden:
        _add_styled_paragraph(doc, "Grondopbouw langs het tracé:", bold=True, space_after=4)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = table.rows[0].cells
        hdr[0].text = "Afstand (m)"
        hdr[1].text = "NAP (m)"
        hdr[2].text = "Grondtype"
        hdr[3].text = "GWS (m)"
        for cell in hdr:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        for d in sorted(boring.doorsneden, key=lambda d: d.volgorde):
            row = table.add_row().cells
            row[0].text = f"{d.afstand_m:.0f}"
            row[1].text = f"{d.NAP_m:+.1f}"
            row[2].text = d.grondtype or "—"
            row[3].text = f"{d.GWS_m:+.1f}" if d.GWS_m else "—"
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
        doc.add_paragraph()

    # Maaiveld uit platform
    if boring.maaiveld_override:
        mv = boring.maaiveld_override
        _add_styled_paragraph(
            doc,
            f"Maaiveld intrede: {mv.MVin_NAP_m:+.2f} m NAP (bron: {mv.MVin_bron or 'handmatig'}). "
            f"Maaiveld uittrede: {mv.MVuit_NAP_m:+.2f} m NAP (bron: {mv.MVuit_bron or 'handmatig'}).",
            space_after=6,
        )

    # 4.1 Sterkteberekeningen
    _add_heading(doc, "4.1.  Sterkteberekeningen", level=2)
    _add_styled_paragraph(
        doc,
        "Om aan te tonen dat de gekozen buizen sterk genoeg zijn om tijdens de aanleg fase ingetrokken "
        "te worden en sterk genoeg zijn om optredende druk van grond- en verkeersbelasting tijdens de "
        "gebruiksfase te weerstaan, zijn de benodigde sterkte berekeningen gemaakt.",
        space_after=6,
    )
    _add_styled_paragraph(
        doc,
        "[Resultaten sterkteberekening invullen - spanning intrekfase en toelaatbare spanning]",
        space_after=6,
        color=RGBColor(0xCC, 0x00, 0x00),
    )

    # 4.2 Intrekkracht
    _add_heading(doc, "4.2.  Berekening benodigde intrekkracht", level=2)
    if boring.berekening and boring.berekening.Ttot_N:
        ton = boring.berekening.Ttot_N / 9810
        _add_styled_paragraph(
            doc,
            f"De benodigde kracht om de HPE buis in te trekken is in par. 5.7 berekend op "
            f"{boring.berekening.Ttot_N:,.0f} N (+/-{ton:,.2f} ton).",
            space_after=12,
        )
    else:
        _add_styled_paragraph(
            doc,
            "[Intrekkracht invullen vanuit Sigma berekening]",
            space_after=12,
            color=RGBColor(0xCC, 0x00, 0x00),
        )

    # 4.3 Boorspoeldruk
    _add_heading(doc, "4.3.  Boorspoeldrukberekeningen", level=2)
    _add_styled_paragraph(
        doc,
        "Er zijn boorspoeldrukberekeningen gemaakt van pilot-, ruim- en intrekfase van de boring.",
        space_after=6,
    )
    _add_styled_paragraph(
        doc,
        "Conclusie: Slechts in de laatste meter van de boring komt de boorspoeldruk boven de "
        "toelaatbare druk uit. Dit ligt in het boorgat, dus is geen sprake van blow-out.",
        space_after=12,
    )

    doc.add_page_break()

    # ── HOOFDSTUK 5: KWEL ───────────────────────────────────────────────

    _add_heading(doc, "5.  Kwel", level=1)
    _add_styled_paragraph(
        doc,
        "Bij werkzaamheden in de (onder)grond in de buurt van waterkeringen en laaggelegen gronden, "
        "dient rekening te worden gehouden met de kans op kwelwater.",
        space_after=12,
    )

    # 5.1 Oorzaak
    _add_heading(doc, "5.1.  Oorzaak", level=2)
    _add_styled_paragraph(
        doc,
        "Tijdens het boorproces wordt grond verwijderd en wordt de oversnijding tussen het boorgat en "
        "de aangebrachte buis opgevuld door de boorspoeling. De grondspanning rondom het boorgat zal "
        "hierdoor veranderen. Nadat de gestuurde boring gereed is, wordt er een nieuw spanningsevenwicht "
        "gevormd tussen de achtergebleven boorspoeling en de grond er omheen. Dit gebeurt doordat het "
        "water langzaam uit de boorspoeling wordt geperst. Daarnaast kan met name in situaties met "
        "relatief zout grondwater de bentoniet na verloop van tijd gaan uitvlokken, waardoor zelfs holle "
        "ruimten in het boorgat ontstaan. Door het veranderen van grondspanning of het ontstaan van holle "
        "ruimten kan grondwater (kwel) gaan stromen. Een kwelstroom kan optreden bij een "
        "waterstandsverschil tussen het in- en uittredepunt. Daarbij moeten niet alleen het open waterpeil, "
        "polderpeilen en de freatische grondwaterstand worden beschouwd, maar ook de stijghoogte "
        "(potentiaal) van het diepe grondwater.",
        space_after=12,
    )

    # 5.2 Risico's en maatregelen
    _add_heading(doc, "5.2.  Risico's en maatregelen", level=2)
    _add_styled_paragraph(
        doc,
        "Het onverwachts ontstaan van kwel zorgt voor overlast en kan bovendien de werking van "
        "waterkeringen negatief beïnvloeden. In geval van twijfel kan er met een kwelwegberekening "
        "worden getoetst of er een kans is op kwel. Vanwege de geringe kosten en het risico van kwel "
        "adviseren wij in geval van twijfel altijd een kwelscherm met een kleikist te plaatsen. "
        "Hierdoor wordt een kwelstroom geblokkeerd indien deze onverwachts toch ontstaat.",
        space_after=12,
    )

    # 5.3 Maatregelen op dit project
    _add_heading(doc, "5.3.  Maatregelen op dit project", level=2)

    if gebruik_ai:
        try:
            from app.ai_assist.werkplan_teksten import genereer_kwel_beoordeling
            kwel_tekst = genereer_kwel_beoordeling(order, boring)
            if kwel_tekst:
                for alinea in kwel_tekst.split("\n\n"):
                    if alinea.strip():
                        _add_styled_paragraph(doc, alinea.strip(), space_after=6)
            else:
                raise ValueError("Lege AI response")
        except Exception:
            # Fallback
            _generate_kwel_template(doc, kwel_gebied)
    else:
        _generate_kwel_template(doc, kwel_gebied)

    doc.add_page_break()

    # ── HOOFDSTUK 6: UITVOERING ─────────────────────────────────────────

    _add_heading(doc, "6.  Uitvoering", level=1)

    # 6.1 In te zetten boormachine
    _add_heading(doc, "6.1   In te zetten boormachine", level=2)
    _add_styled_paragraph(
        doc,
        "De Certificatieregeling Kabelinfrastructuur en Buizenlegbedrijven (CKB) kent een onderverdeling "
        "van boormachines in 3 processen:",
        space_after=6,
    )
    _add_styled_paragraph(doc, "ST-A: Kleine gestuurde boringen (maximaal 9 ton),", space_after=2)
    _add_styled_paragraph(doc, "ST-B: Grote gestuurde boringen (van 9 tot 39 ton),", space_after=2)
    _add_styled_paragraph(doc, "ST-C: Zeer grote gestuurde boringen (van 40 tot 149 ton),", space_after=2)
    _add_styled_paragraph(doc, "ST-D: Extra grote gestuurde boringen (meer dan 150 ton).", space_after=6)

    ckb_cat, ckb_desc = _ckb_categorie(boring)
    _add_styled_paragraph(
        doc,
        f"Dit project is conform de CKB-regeling gecategoriseerd als een {ckb_desc.lower()}, welke "
        f"uitgevoerd kan worden met een {ckb_desc.split('(')[0].strip().lower()} boormachine "
        f"(proces {ckb_cat}). Op dit project is er geen aanleiding om van de CKB regeling af te wijken.",
        space_after=12,
    )

    # 6.1.1 Details boormachine
    _add_heading(doc, "6.1.1   Details boormachine", level=3)
    _add_styled_paragraph(
        doc,
        "[Boormachine details invullen - zie bijlage E]",
        space_after=12,
        color=RGBColor(0xCC, 0x00, 0x00),
    )

    # 6.2 Personeel
    _add_heading(doc, "6.2   Personeel", level=2)
    _add_styled_paragraph(doc, "De boorploeg voor uitvoering van deze boring zal bestaan uit:", space_after=6)

    pers_table = doc.add_table(rows=3, cols=2)
    pers_data = [
        ("1 Boormeester/voorman:",
         "Algehele leiding over het project, bediening van de boormachine zorg "
         "voor een goede uitvoering van de pilotboring en de maatvoering"),
        ("1 Assistent Boormeester:",
         "Verzorgen van de boorspoeling en de metingen tijdens de pilotfase"),
        ("1 à 2 medewerkers:",
         "Mengen van de Bentoniet, koppelen en invoeren van de leiding, afvoeren van de Bentoniet"),
    ]
    for i, (rol, desc) in enumerate(pers_data):
        pers_table.rows[i].cells[0].text = rol
        pers_table.rows[i].cells[1].text = desc
        for cell in pers_table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"

    doc.add_paragraph()

    # 6.3 Tijdsplanning
    _add_heading(doc, "6.3   Tijdsplanning", level=2)
    _add_styled_paragraph(doc, "De planning van de werkzaamheden ziet er als volgt uit:", space_after=6)

    _add_styled_paragraph(doc, "Dag 1:", bold=True, space_after=4)
    dag1_items = [
        ("KLIC melding controleren:", "0,5 uur"),
        ("Graven boorgaten:", "2 uur"),
    ]
    for item, uur in dag1_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{item}\t\t{uur}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    doc.add_paragraph()
    _add_styled_paragraph(doc, "Dag 2:", bold=True, space_after=4)
    dag2_items = [
        ("Aanvoer materieel, opstellen:", "2 uur"),
        ("KLIC melding controleren:", "0,5 uur"),
        ("Pilotboring:", "2,5 uur"),
        ("Ruimgang:", "1 uur"),
        ("Centreren en intrekken van de buis:", "1 uur"),
        ("Werklocatie herstellen:", "1 uur"),
        ("Materieel afvoeren:", "1 uur"),
    ]
    for item, uur in dag2_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{item}\t\t{uur}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    doc.add_paragraph()

    # 6.4 Afwijkingen
    _add_heading(doc, "6.4   Afwijkingen", level=2)
    _add_styled_paragraph(
        doc,
        "Het is mogelijk dat de pilotboring door bepaalde omstandigheden afwijkt van het vooraf geplande "
        "boortracé.",
        space_after=6,
    )
    _add_styled_paragraph(
        doc,
        "De boormeester en assistent-boormeester houden continu en nauwlettend het boorproces in de "
        "gaten en eventuele afwijkingen worden geregistreerd.",
        space_after=6,
    )
    _add_styled_paragraph(
        doc,
        "Wanneer de boring de 'maximale toegestane afwijking' overschrijdt zal direct contact worden "
        "opgenomen met de opdrachtgever. Na overleg met de opdrachtgever wordt bepaald of er passende "
        "maatregelen dienen te worden genomen om de gestuurde boring te kunnen vervolgen of dat de "
        "gestuurde boring dient te worden beëindigd.",
        space_after=6,
    )
    _add_styled_paragraph(
        doc,
        "De 'maximale toegestane uitvoeringsafwijkingen' van de boorlijn zijn weergegeven in de "
        "onderstaande tabel:",
        space_after=6,
    )

    # Afwijkingen tabel
    afw_table = doc.add_table(rows=6, cols=2)
    afw_table.style = "Table Grid"
    afw_data = [
        ("Richting:", "Maximale uitvoeringsafwijking:"),
        ("Verticaal", "+0,25 / -0,25 m¹"),
        ("Horizontaal:", ""),
        ("  - in lengterichting; t.p.v. uittredepunt", "+0,5 / -0,5 m¹"),
        ("  - in dwarsrichting; t.p.v. uittredepunt", "+0,5 / -0,5 m¹"),
        ("Bochtstralen", "< 10%"),
    ]
    for i, (label, val) in enumerate(afw_data):
        afw_table.rows[i].cells[0].text = label
        afw_table.rows[i].cells[1].text = val
        for cell in afw_table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
                    if i == 0:
                        run.bold = True

    doc.add_paragraph()

    # 6.4 Inspectie na intrekken
    _add_heading(doc, "6.5   Inspectie na intrekken", level=2)
    _add_styled_paragraph(
        doc,
        "Direct na het afronden van de intrekfase dient de polyetheen leiding bij het intredepunt van de "
        "pilotboring (het uittredepunt voor de ingetrokken leiding) het vrijgekomen stuk leiding "
        "geïnspecteerd te worden. De getrokken leiding moet voldoen aan de in NEN7244-2 gestelde eisen.",
        space_after=12,
    )

    doc.add_page_break()

    # ── BIJLAGEN (placeholders) ─────────────────────────────────────────

    bijlage_titels = [
        "Bijlage A: Tekening van de gestuurde boring",
        "Bijlage B: Berekeningen",
        "Bijlage C: Geotechnische bodemgegevens",
        "Bijlage D: Procesbeschrijving techniek van gestuurd boren",
        "Bijlage E: De boormachine",
        "Bijlage F: Het boorkop-volgsysteem",
        "Bijlage G: Bentoniet",
    ]

    for titel in bijlage_titels:
        _add_heading(doc, titel, level=1)
        doc.add_paragraph()
        _add_styled_paragraph(
            doc, "Zie volgende pagina",
            bold=True, space_after=12,
        )
        doc.add_page_break()

    # ── Footer instellen ────────────────────────────────────────────────
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Opdrachtnummer + locatie
    boring_naam = boring.naam or ""
    run = fp.add_run(
        f"{order.ordernummer} Werkplan - {locatie}"
    )
    run.font.size = Pt(8)
    run.font.name = "Calibri"

    # Tab + datum
    run = fp.add_run(f"\t{vandaag}")
    run.font.size = Pt(8)
    run.font.name = "Calibri"

    # ── Opslaan naar bytes ──────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
